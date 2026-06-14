from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Title Page ──
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Layover Explorer')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(30, 58, 138)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Project Description Document')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Exercise 3 — Spring 2026\n').font.size = Pt(12)
info.add_run('Ofek Cohen').font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════
# 1. GENERAL EXPLANATION
# ══════════════════════════════════════════════
doc.add_heading('1. General Explanation', level=1)

doc.add_paragraph(
    'Layover Explorer is a flight search engine designed for budget-conscious travelers '
    'who are willing to take connecting flights but want to make the most of their layover time. '
    'Instead of treating layovers as wasted hours stuck at an airport, Layover Explorer transforms '
    'them into mini city-exploration opportunities.'
)

doc.add_paragraph(
    'The core idea: A traveler wants to fly from Tel Aviv to New York. A direct ticket costs $1,000, '
    'but a connecting flight through London with a 7-hour layover costs only $600. Rather than sitting '
    'at Heathrow, the system detects that 7 hours is enough time to exit the airport, visit the British '
    'Museum (45-minute transit), explore for ~2 hours, eat lunch, and return to the airport with a '
    '2-hour security buffer. The system builds this plan automatically based on the user\'s interests.'
)

doc.add_heading('Key Features', level=2)

features = [
    ('Interest-Based Matching', 'Users specify their interests (arts & museums, nightlife, food & dining). '
     'The recommendation engine scores layover cities based on matching Points of Interest (POIs).'),
    ('Itinerary Planning', 'For each viable layover, the system calculates a time-breakdown: '
     'deboarding time → transit to attraction → exploration/eating time → return transit → security buffer.'),
    ('Flight Comparison', 'Results are ranked by a Match Score combining price, layover quality, '
     'and interest relevance, helping users find the best value.'),
    ('User Profiles & Preferences', 'Registered users can save their interest preferences, view search history, '
     'and bookmark flights for later.'),
    ('Admin Dashboard', 'Administrators can monitor user activity, view search statistics, '
     'manage POI data, and oversee system health.'),
    ('Session-Based Experience', 'Recent searches are stored in the session for quick re-access. '
     'Users can save flights to a "shortlist" before logging in, which persists after authentication.'),
]

for title_text, desc in features:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('Technology Stack', level=2)

tech = [
    ('Backend', 'Java 17+, Spring Boot 3, Spring MVC, Spring Security, Spring Data JPA'),
    ('View Engine', 'Thymeleaf (server-side rendering) for all pages except the flight search'),
    ('SPA Page', 'React (Vite) for the interactive flight search & results page'),
    ('Database', 'MySQL (database name: ex4)'),
    ('Build Tool', 'Maven with Maven Wrapper (mvnw)'),
    ('Security', 'Spring Security with BCrypt password encoding, role-based authorization (USER / ADMIN)'),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Component'
hdr[1].text = 'Technology'
for comp, tech_text in tech:
    row = table.add_row().cells
    row[0].text = comp
    row[1].text = tech_text

doc.add_page_break()

# ══════════════════════════════════════════════
# 2. MAIN PAGES
# ══════════════════════════════════════════════
doc.add_heading('2. Main Pages and Their Goals', level=1)

doc.add_paragraph(
    'The application consists of at least 6 major pages. Five are server-side rendered using '
    'Thymeleaf, and one (the Flight Search page) is implemented as a Single Page Application using React.'
)

pages = [
    {
        'name': 'Landing / Home Page',
        'route': '/',
        'type': 'Thymeleaf',
        'goal': 'The entry point of the application. Displays a hero section explaining the concept of '
                'Layover Explorer, a call-to-action to search flights or register, and highlights of '
                'popular layover cities. If the user is logged in, it shows a personalized greeting '
                'and quick-access links to recent searches (pulled from session).',
        'features': [
            'Hero banner with tagline and search shortcut',
            'Featured layover destinations carousel',
            'Quick access to recent searches (session-based)',
            'Login/Register buttons (or profile link if authenticated)',
        ]
    },
    {
        'name': 'Login Page',
        'route': '/login',
        'type': 'Thymeleaf (Spring Security)',
        'goal': 'Handles user authentication via Spring Security form-based login. Displays error '
                'messages for invalid credentials. Redirects to the originally requested page after '
                'successful login. Provides a link to the registration page.',
        'features': [
            'Username/password form integrated with Spring Security',
            'Error feedback for failed login attempts',
            'Link to registration page',
            'CSRF token protection',
        ]
    },
    {
        'name': 'Registration Page',
        'route': '/register',
        'type': 'Thymeleaf',
        'goal': 'Allows new users to create an account. Collects username, email, password (with '
                'confirmation), and initial interest preferences. Server-side validation ensures '
                'unique usernames/emails and password strength. On success, auto-logs the user in.',
        'features': [
            'Registration form with server-side validation',
            'Interest preference selection during sign-up (arts, nightlife, food)',
            'Password strength validation',
            'Duplicate username/email detection',
        ]
    },
    {
        'name': 'User Profile / Dashboard',
        'route': '/profile',
        'type': 'Thymeleaf',
        'goal': 'A personal dashboard for authenticated users. Displays and allows editing of '
                'interest preferences, shows search history with clickable re-search links, and '
                'lists saved/bookmarked flights. Users can manage their account settings.',
        'features': [
            'Edit interest preferences (persisted via JPA)',
            'View and clear search history',
            'View and manage saved flights (with remove option)',
            'Account settings (change password, update email)',
        ]
    },
    {
        'name': 'Flight Search & Results (SPA)',
        'route': '/search',
        'type': 'React (Single Page Application)',
        'goal': 'The core interactive page of the application. Users enter origin, destination, '
                'and interest. The React frontend calls REST API endpoints to fetch scored flight '
                'results with layover itineraries. Results display the full layover plan including '
                'deboarding time, transit, exploration time, and security buffer for each POI.',
        'features': [
            'Real-time search with loading states',
            'Interactive flight result cards with expandable itineraries',
            'Match Score and star ratings for each layover',
            'Save/bookmark flights (via REST API, persisted to DB for logged-in users)',
            'POI itinerary breakdown with timing for each activity',
        ]
    },
    {
        'name': 'Admin Dashboard',
        'route': '/admin',
        'type': 'Thymeleaf (ADMIN role required)',
        'goal': 'A backend management panel accessible only to users with the ADMIN role. '
                'Provides oversight of the system including user management, search statistics, '
                'POI data management, and system configuration.',
        'features': [
            'User list with role management (promote/demote, disable accounts)',
            'Search statistics (most searched routes, popular interests)',
            'POI management (add/edit/delete Points of Interest)',
            'System monitoring (total users, total searches, recent activity log)',
        ]
    },
]

for page in pages:
    doc.add_heading(page['name'], level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('Route: ')
    run.bold = True
    p.add_run(page['route'])
    
    p = doc.add_paragraph()
    run = p.add_run('Rendering: ')
    run.bold = True
    p.add_run(page['type'])
    
    p = doc.add_paragraph()
    run = p.add_run('Goal: ')
    run.bold = True
    p.add_run(page['goal'])
    
    doc.add_paragraph('Key Features:', style='List Bullet')
    for feat in page['features']:
        doc.add_paragraph(feat, style='List Bullet 2')

doc.add_page_break()

# ══════════════════════════════════════════════
# 3. DATABASE BEANS
# ══════════════════════════════════════════════
doc.add_heading('3. Database Entities (JPA Beans)', level=1)

doc.add_paragraph(
    'The application uses Spring Data JPA with a MySQL database named "ex4". '
    'There are 5 entity classes (tables) with the following relations:'
)

# ── Entity Relationship Summary ──
doc.add_heading('Entity Relationship Overview', level=2)

relations = [
    'User (1) ←→ (1) UserPreference — One-to-One',
    'User (1) ←→ (N) SearchHistory — One-to-Many',
    'User (1) ←→ (N) SavedFlight — One-to-Many',
    'PointOfInterest — Standalone reference table (no FK to User)',
]

for r in relations:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph()

# ── Diagram description ──
p = doc.add_paragraph()
run = p.add_run('Relationship Diagram:')
run.bold = True

doc.add_paragraph(
    '┌──────────┐     1:1     ┌─────────────────┐\n'
    '│   User   │────────────▶│ UserPreference   │\n'
    '└──────────┘             └─────────────────┘\n'
    '     │\n'
    '     │  1:N\n'
    '     ├──────────────────▶┌─────────────────┐\n'
    '     │                   │ SearchHistory    │\n'
    '     │                   └─────────────────┘\n'
    '     │  1:N\n'
    '     └──────────────────▶┌─────────────────┐\n'
    '                         │ SavedFlight      │\n'
    '                         └─────────────────┘\n'
    '\n'
    '┌─────────────────┐\n'
    '│ PointOfInterest  │  (standalone)\n'
    '└─────────────────┘'
).style = 'No Spacing'

doc.add_paragraph()

# ── Entity Details ──
entities = [
    {
        'name': 'User',
        'table': 'users',
        'desc': 'Represents a registered user of the application. Stores authentication credentials '
                'and role information. Serves as the parent entity for preferences, search history, and saved flights.',
        'fields': [
            ('id', 'Long', 'Primary key, auto-generated'),
            ('username', 'String', 'Unique, not null. Used for login'),
            ('email', 'String', 'Unique, not null. User email address'),
            ('password', 'String', 'Not null. BCrypt-encoded password hash'),
            ('role', 'String (Enum)', 'USER or ADMIN. Defaults to USER'),
            ('createdAt', 'LocalDateTime', 'Timestamp of account creation'),
            ('enabled', 'Boolean', 'Whether the account is active. Default: true'),
        ],
        'relations': [
            '@OneToOne(mappedBy = "user", cascade = ALL) → UserPreference',
            '@OneToMany(mappedBy = "user", cascade = ALL) → List<SearchHistory>',
            '@OneToMany(mappedBy = "user", cascade = ALL) → List<SavedFlight>',
        ]
    },
    {
        'name': 'UserPreference',
        'table': 'user_preferences',
        'desc': 'Stores user-specific preferences for the recommendation engine. Each user has exactly one '
                'preference record. Interests determine how layover cities are scored and ranked.',
        'fields': [
            ('id', 'Long', 'Primary key, auto-generated'),
            ('user', 'User', 'Foreign key to User (owning side of @OneToOne)'),
            ('interestArts', 'Boolean', 'User interested in arts & museums'),
            ('interestNightlife', 'Boolean', 'User interested in nightlife & clubs'),
            ('interestFood', 'Boolean', 'User interested in food & dining'),
            ('minLayoverHours', 'Integer', 'Minimum acceptable layover duration. Default: 4'),
            ('maxLayoverHours', 'Integer', 'Maximum acceptable layover duration. Default: 24'),
        ],
        'relations': [
            '@OneToOne → User (with @JoinColumn(name = "user_id"))',
        ]
    },
    {
        'name': 'SearchHistory',
        'table': 'search_history',
        'desc': 'Records every flight search performed by a registered user. Used to display recent '
                'searches on the profile page and to power the session-based "quick access" feature. '
                'Also feeds the admin dashboard statistics.',
        'fields': [
            ('id', 'Long', 'Primary key, auto-generated'),
            ('user', 'User', 'Foreign key to User (@ManyToOne)'),
            ('origin', 'String', 'IATA code of departure airport (e.g., TLV)'),
            ('destination', 'String', 'IATA code of destination airport (e.g., JFK)'),
            ('interest', 'String', 'Interest category used in this search (arts/nightlife/food/any)'),
            ('searchDate', 'LocalDateTime', 'Timestamp of when the search was performed'),
            ('resultsCount', 'Integer', 'Number of flight results returned'),
        ],
        'relations': [
            '@ManyToOne → User (with @JoinColumn(name = "user_id"))',
        ]
    },
    {
        'name': 'SavedFlight',
        'table': 'saved_flights',
        'desc': 'Represents a flight that a user has bookmarked/saved for later review. Stores a '
                'snapshot of the flight details and its recommendation score at the time of saving. '
                'Users can manage their saved flights from their profile page.',
        'fields': [
            ('id', 'Long', 'Primary key, auto-generated'),
            ('user', 'User', 'Foreign key to User (@ManyToOne)'),
            ('flightId', 'String', 'External flight identifier'),
            ('airline', 'String', 'Airline name (e.g., British Airways)'),
            ('price', 'Double', 'Flight price in USD'),
            ('origin', 'String', 'IATA code of departure airport'),
            ('destination', 'String', 'IATA code of final destination'),
            ('layoverCity', 'String', 'IATA code of the layover city'),
            ('layoverDurationHours', 'Double', 'Duration of the layover in hours'),
            ('recommendationScore', 'Double', 'Match score at time of saving'),
            ('savedAt', 'LocalDateTime', 'Timestamp of when the flight was saved'),
        ],
        'relations': [
            '@ManyToOne → User (with @JoinColumn(name = "user_id"))',
        ]
    },
    {
        'name': 'PointOfInterest',
        'table': 'points_of_interest',
        'desc': 'A reference table containing attractions, restaurants, and venues near major airports. '
                'Used by the Recommendation Engine to match layover cities with user interests. '
                'Managed by admins via the admin dashboard. This is a standalone entity with no '
                'foreign key relations to User.',
        'fields': [
            ('id', 'Long', 'Primary key, auto-generated'),
            ('name', 'String', 'Name of the point of interest (e.g., The Louvre)'),
            ('airportCode', 'String', 'IATA code of the nearest airport (e.g., CDG)'),
            ('city', 'String', 'City name (e.g., Paris)'),
            ('category', 'String', 'Interest category: arts, nightlife, or food'),
            ('rating', 'Double', 'Quality rating (1.0–5.0)'),
            ('transitTimeMinutes', 'Integer', 'Estimated one-way transit time from airport in minutes'),
            ('description', 'String', 'Brief description of the attraction'),
        ],
        'relations': [
            'None — standalone reference data table',
        ]
    },
]

for entity in entities:
    doc.add_heading(f'{entity["name"]} Entity', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('Table name: ')
    run.bold = True
    p.add_run(entity['table'])
    
    doc.add_paragraph(entity['desc'])
    
    # Fields table
    doc.add_heading('Fields', level=3)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Field'
    hdr[1].text = 'Type'
    hdr[2].text = 'Description'
    for field_name, field_type, field_desc in entity['fields']:
        row = table.add_row().cells
        row[0].text = field_name
        row[1].text = field_type
        row[2].text = field_desc
    
    # Relations
    doc.add_heading('JPA Relations', level=3)
    for rel in entity['relations']:
        doc.add_paragraph(rel, style='List Bullet')
    
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════
# 4. SESSION USAGE
# ══════════════════════════════════════════════
doc.add_heading('4. Session Usage', level=1)

doc.add_paragraph(
    'HTTP sessions are used throughout the application to enhance user experience:'
)

session_uses = [
    ('Recent Searches', 'The last 5 searches (origin, destination, interest) are stored in the '
     'session and displayed on the home page and search page for quick re-access.'),
    ('Flight Shortlist', 'Before logging in, users can "save for later" flights. These are stored '
     'in the session. Upon login, the session-saved flights are automatically persisted to the '
     'database under the user\'s SavedFlight records.'),
    ('Search Context', 'The current search parameters are stored in session so that navigating '
     'away from the search page and returning preserves the user\'s last search state.'),
]

for title_text, desc in session_uses:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

# ══════════════════════════════════════════════
# 5. SPRING SECURITY
# ══════════════════════════════════════════════
doc.add_heading('5. Security Architecture', level=1)

doc.add_paragraph(
    'The application uses Spring Security for authentication and authorization:'
)

security_items = [
    'Form-based login with custom Thymeleaf login page',
    'BCrypt password encoding for all stored passwords',
    'Two roles: USER (default) and ADMIN',
    'Route protection: /profile/**, /admin/** require authentication; /admin/** requires ADMIN role',
    'CSRF protection enabled on all forms',
    'Auto-creation of default admin account on startup if none exists (username: admin, password: admin123)',
    'Public access to: /, /login, /register, /search, /api/** (REST endpoints for React SPA)',
]

for item in security_items:
    doc.add_paragraph(item, style='List Bullet')

# ── Save ──
output_path = r'c:\Users\cofek\WebstormProjects\layover-explorer\Layover_Explorer_Project_Description.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
